from __future__ import annotations

from html import escape
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from document_processing_pipeline import bootstrap as dependency_bootstrap
from document_processing_pipeline.io_helpers import ordered_blocks as _ordered_blocks



def assemble_docx(run_dir: str | Path) -> Path:
    run_path = Path(run_dir)
    output_path = run_path / "output.docx"

    try:
        dependency_bootstrap.ensure_python_dependency("python_docx")
        import docx  # type: ignore
    except (ImportError, dependency_bootstrap.DependencyBootstrapError):
        from document_processing_pipeline.models import RichDocument

        transformed = run_path / "rich_ir.transformed.json"
        document = RichDocument.load_json(transformed if transformed.exists() else run_path / "rich_ir.json")
        _write_basic_docx(document, output_path)
        return output_path

    from document_processing_pipeline.models import RichDocument

    transformed = run_path / "rich_ir.transformed.json"
    document = RichDocument.load_json(transformed if transformed.exists() else run_path / "rich_ir.json")
    doc = docx.Document()
    for block in _ordered_blocks(document):
        if block.block_type in {"title", "heading"}:
            level = int(block.metadata.get("heading_level", 1))
            doc.add_heading(block.text, level=min(level, 9))
        else:
            doc.add_paragraph(block.text)
    doc.save(str(output_path))
    return output_path


def _paragraph_xml(text: str, style: str | None = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return f"<w:p>{style_xml}<w:r><w:t xml:space=\"preserve\">{escape(text)}</w:t></w:r></w:p>"


def _write_basic_docx(document, output_path: Path) -> None:
    body_parts: list[str] = []
    for block in _ordered_blocks(document):
        text = block.text or ""
        style = None
        if block.block_type == "title":
            level = int(block.metadata.get("heading_level", 1))
            style = "Heading1" if level <= 1 else "Heading2"
        elif block.block_type == "heading":
            level = int(block.metadata.get("heading_level", 2))
            style = "Heading1" if level <= 1 else "Heading2"
        body_parts.append(_paragraph_xml(text, style=style))

    document_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:document xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas" '
        'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" '
        'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word" '
        'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml" '
        'xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup" '
        'xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk" '
        'xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" '
        'mc:Ignorable="w14 wp14">'
        f"<w:body>{''.join(body_parts)}<w:sectPr><w:pgSz w:w=\"11906\" w:h=\"16838\"/><w:pgMar w:top=\"1440\" w:right=\"1440\" w:bottom=\"1440\" w:left=\"1440\" w:header=\"708\" w:footer=\"708\" w:gutter=\"0\"/></w:sectPr></w:body>"
        "</w:document>"
    )
    styles_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/></w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/></w:style>'
        '<w:style w:type="paragraph" w:styleId="Heading2"><w:name w:val="heading 2"/></w:style>'
        "</w:styles>"
    )

    with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            '<Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>'
            "</Types>",
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>'
            "</Relationships>",
        )
        archive.writestr(
            "word/_rels/document.xml.rels",
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/>'
            "</Relationships>",
        )
        archive.writestr("word/document.xml", document_xml)
        archive.writestr("word/styles.xml", styles_xml)
